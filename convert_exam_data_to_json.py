#!/usr/bin/env python3
from __future__ import annotations

import csv
import json
import os
import re
import shutil
import sys
import zipfile
from collections import Counter, defaultdict
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import fitz
import openpyxl
import pdfplumber
from docx import Document


ROOT = Path(__file__).resolve().parent
WORK_DIR = ROOT / "_work_extracted"
ASSETS_DIR = ROOT / "assets"
OUTPUT_JSON = ROOT / "output.json"
ERRORS_JSON = ROOT / "errors.json"
SUMMARY_JSON = ROOT / "summary.json"

QUESTION_PATTERN = re.compile(r"^(\d{1,3})[.．、](.*)$")
QUESTION_PREFIX_ONLY_PATTERN = re.compile(r"^(\d{1,3})[.．、]\s*$")
OPTION_PATTERN = re.compile(r"^([A-DＡ-Ｄ])[.．、](.*)$")
OPTION_PREFIX_ONLY_PATTERN = re.compile(r"^([A-DＡ-Ｄ])[.．、]\s*$")
YEAR_ROUND_PATTERN = re.compile(r"(\d{2,3})\s*年\s*(?:第)?\s*([12一二])\s*次")
SUBJECT_PATTERN = re.compile(r"科目名稱[:：]\s*([^\n]+)")
CLASS_NAME_PATTERN = re.compile(r"類科名稱[:：]\s*([^\n]+)")
QUESTION_KEY = Tuple[int, str, str, int]

FULLWIDTH_TRANS = str.maketrans(
    {
        "Ａ": "A",
        "Ｂ": "B",
        "Ｃ": "C",
        "Ｄ": "D",
        "＃": "#",
        "０": "0",
        "１": "1",
        "２": "2",
        "３": "3",
        "４": "4",
        "５": "5",
        "６": "6",
        "７": "7",
        "８": "8",
        "９": "9",
        "（": "(",
        "）": ")",
        "，": ",",
        "：": ":",
        "　": " ",
    }
)

IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}
TEXT_SUFFIXES = {".txt", ".csv"}
QUESTION_SOURCE = "考選部歷屆試題"


def normalize_text(value: str) -> str:
    return value.replace("\xa0", " ").replace("\u3000", " ").replace("\r", "").strip()


def normalize_ascii(value: str) -> str:
    return normalize_text(value).translate(FULLWIDTH_TRANS)


def normalize_exam_round(raw: str) -> Optional[str]:
    raw = normalize_ascii(raw)
    if "1" in raw or "一" in raw:
        return "第一次"
    if "2" in raw or "二" in raw:
        return "第二次"
    return None


def round_number_from_exam_round(exam_round: str) -> str:
    return "1" if exam_round == "第一次" else "2"


def build_question_id(year: int, exam_round: str, subject: str, question_number: int) -> str:
    return f"{year}-{round_number_from_exam_round(exam_round)}-{subject}-{question_number:03d}"


def append_text(target: str, addition: str) -> str:
    addition = normalize_text(addition)
    if not addition:
        return target
    if not target:
        return addition
    return f"{target}\n{addition}"


def detect_file_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".zip":
        return "zip"
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix in {".xlsx", ".xlsm"}:
        return "xlsx"
    if suffix in TEXT_SUFFIXES:
        return "text"
    if suffix in IMAGE_SUFFIXES:
        return "image"
    return "unsupported"


def detect_document_role(path: Path) -> str:
    name = normalize_ascii(path.name.upper())
    if "_MOD" in name or any(keyword in name for keyword in ["更正", "疑義", "MOD"]):
        return "corrected_answer"
    if "_ANS" in name or "答案" in path.name:
        return "answer"
    return "question"


def resolve_input_folder(arg_path: str) -> Path:
    requested = Path(arg_path)
    if requested.is_absolute() and requested.exists():
        return requested

    candidates = [
        ROOT / requested,
        ROOT / "src" / "data" / requested.name,
        ROOT / "src" / "data" / "questions_original_data",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    raise FileNotFoundError(f"找不到輸入資料夾：{arg_path}")


def scan_input_folder(input_dir: Path) -> List[Path]:
    return sorted(path for path in input_dir.rglob("*") if path.is_file())


def open_zip_with_fallback(path: Path) -> zipfile.ZipFile:
    attempts = ["cp950", "big5", None, "utf-8"]
    for encoding in attempts:
        try:
            if encoding is None:
                return zipfile.ZipFile(path)
            return zipfile.ZipFile(path, metadata_encoding=encoding)
        except Exception:
            continue
    raise zipfile.BadZipFile(f"無法開啟 zip：{path}")


def extract_zip_files(zip_paths: Iterable[Path]) -> Tuple[List[Path], int, List[dict]]:
    if WORK_DIR.exists():
        shutil.rmtree(WORK_DIR)
    WORK_DIR.mkdir(parents=True, exist_ok=True)

    extracted_zip_paths: List[Path] = []
    errors: List[dict] = []
    processed = set()
    pending = list(zip_paths)
    zip_count = 0

    while pending:
        zip_path = pending.pop(0)
        zip_path = zip_path.resolve()
        if zip_path in processed:
            continue
        processed.add(zip_path)
        zip_count += 1

        target_dir = WORK_DIR / zip_path.stem
        target_dir.mkdir(parents=True, exist_ok=True)

        try:
          with open_zip_with_fallback(zip_path) as archive:
                archive.extractall(target_dir)
        except Exception as exc:
            errors.append(
                {
                    "type": "parse_failed",
                    "file": str(zip_path),
                    "message": f"zip 解壓失敗：{exc}",
                }
            )
            continue

        extracted_zip_paths.append(target_dir)
        for nested_zip in target_dir.rglob("*.zip"):
            if nested_zip.resolve() not in processed:
                pending.append(nested_zip)

    return extracted_zip_paths, zip_count, errors


def parse_pdf_metadata(doc: fitz.Document) -> Tuple[Optional[int], Optional[str], Optional[str], Optional[str]]:
    text = "\n".join(page.get_text("text") for page in doc[:2])
    year_match = YEAR_ROUND_PATTERN.search(text)
    subject_match = SUBJECT_PATTERN.search(text)
    class_name_match = CLASS_NAME_PATTERN.search(text)
    year = int(year_match.group(1)) if year_match else None
    exam_round = normalize_exam_round(year_match.group(2)) if year_match else None
    subject = None
    if subject_match:
        subject = subject_match.group(1).split("(試題代號", 1)[0].strip()
    class_name = class_name_match.group(1).strip() if class_name_match else None
    return year, exam_round, subject, class_name


def save_pdf_image(image_bytes: bytes, question_id: str, image_index: int) -> str:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    image_path = ASSETS_DIR / f"{question_id}-{image_index}.png"
    with open(image_path, "wb") as handle:
        handle.write(image_bytes)
    return str(image_path.relative_to(ROOT))


def finalize_question(question: dict, errors: List[dict], source_file: Path) -> Optional[dict]:
    if not question:
        return None
    missing_options = [key for key in ["A", "B", "C", "D"] if not question["options"].get(key)]
    if missing_options:
        errors.append(
            {
                "type": "missing_options",
                "year": question["year"],
                "exam_round": question["exam_round"],
                "subject": question["subject"],
                "question_number": question["question_number"],
                "file": str(source_file),
                "message": f"缺少選項：{','.join(missing_options)}",
            }
        )
        return None
    question["question"] = question["question"].strip()
    for option_key in ["A", "B", "C", "D"]:
        question["options"][option_key] = question["options"][option_key].strip()
    return question


def parse_question_pdf(path: Path, errors: List[dict]) -> List[dict]:
    questions: List[dict] = []
    doc = fitz.open(path)
    year, exam_round, subject, class_name = parse_pdf_metadata(doc)
    if class_name != "醫事放射師":
        return questions
    if year is None or exam_round is None or not subject:
        errors.append(
            {
                "type": "parse_failed",
                "file": str(path),
                "message": "無法從題目 PDF 辨識年份、考試次別或科目",
            }
        )
        return questions

    current_question: Optional[dict] = None
    current_option: Optional[str] = None
    question_counter_images: Dict[int, int] = defaultdict(int)

    def start_question(question_number: int, question_text: str = "") -> None:
        nonlocal current_question, current_option
        if question_number <= 0:
            return
        current_question = {
            "id": build_question_id(year, exam_round, subject, question_number),
            "year": year,
            "exam_round": exam_round,
            "subject": subject,
            "question_number": question_number,
            "question": normalize_text(question_text),
            "options": {"A": "", "B": "", "C": "", "D": ""},
            "explanation": "",
            "source": QUESTION_SOURCE,
            "images": [],
        }
        current_option = None

    for page in doc:
        page_dict = page.get_text("dict")
        events: List[Tuple[float, float, str, object]] = []

        for y_mid, entries in collect_word_rows(page):
            if not entries:
                continue
            x0 = entries[0][0]
            row_text = normalize_text(" ".join(text for _, text in entries))
            if row_text:
                events.append((y_mid, x0, "text", row_text))

        for block in page_dict["blocks"]:
            if block["type"] == 1:
                events.append((block["bbox"][1], block["bbox"][0], "image", block))

        for _, _, event_type, payload in sorted(events, key=lambda item: (item[0], item[1])):
            if event_type == "image":
                block = payload
                if not current_question:
                    continue
                try:
                    question_number = current_question["question_number"]
                    question_counter_images[question_number] += 1
                    question_id = build_question_id(year, exam_round, subject, question_number)
                    relative_path = save_pdf_image(
                        block["image"],
                        question_id,
                        question_counter_images[question_number],
                    )
                    placeholder = f"{{{{image:{relative_path}}}}}"
                    position = "question"
                    if current_option:
                        position = f"options.{current_option}"
                        current_question["options"][current_option] = append_text(
                            current_question["options"][current_option],
                            placeholder,
                        )
                    else:
                        current_question["question"] = append_text(current_question["question"], placeholder)
                    current_question["images"].append(
                        {
                            "path": relative_path,
                            "position": position,
                            "placeholder": placeholder,
                        }
                    )
                except Exception as exc:
                    errors.append(
                        {
                            "type": "image_extract_failed",
                            "year": year,
                            "exam_round": exam_round,
                            "subject": subject,
                            "file": str(path),
                            "message": f"圖片抽取失敗：{exc}",
                        }
                    )
                continue

            line = normalize_text(str(payload))
            if not line or line.startswith("※") or line.startswith("代") or line.startswith("類科名稱") or line.startswith("科目名稱") or line.startswith("考試時間") or line.startswith("座號") or "醫事放射師考試" in line:
                continue
            question_match = QUESTION_PATTERN.match(line)
            question_prefix_only_match = QUESTION_PREFIX_ONLY_PATTERN.match(line)
            option_match = OPTION_PATTERN.match(line)
            option_prefix_only_match = OPTION_PREFIX_ONLY_PATTERN.match(line)

            if question_match:
                finalized = finalize_question(current_question, errors, path)
                if finalized:
                    questions.append(finalized)
                question_number = int(question_match.group(1))
                start_question(question_number, question_match.group(2))
                continue

            if question_prefix_only_match:
                finalized = finalize_question(current_question, errors, path)
                if finalized:
                    questions.append(finalized)
                question_number = int(question_prefix_only_match.group(1))
                start_question(question_number)
                continue

            if option_match and current_question:
                current_option = normalize_ascii(option_match.group(1))
                current_question["options"][current_option] = append_text(
                    current_question["options"][current_option],
                    option_match.group(2),
                )
                continue

            if option_prefix_only_match and current_question:
                current_option = normalize_ascii(option_prefix_only_match.group(1))
                continue

            if current_question:
                if current_option:
                    current_question["options"][current_option] = append_text(
                        current_question["options"][current_option],
                        line,
                    )
                else:
                    current_question["question"] = append_text(current_question["question"], line)

    finalized = finalize_question(current_question, errors, path)
    if finalized:
        questions.append(finalized)
    return questions


def collect_word_rows(page: fitz.Page, tolerance: float = 3.0) -> List[Tuple[float, List[Tuple[float, str]]]]:
    rows: List[Tuple[float, List[Tuple[float, str]]]] = []
    words = sorted(page.get_text("words"), key=lambda item: ((item[1] + item[3]) / 2, item[0]))
    for x0, y0, x1, y1, text, *_ in words:
        y_mid = (y0 + y1) / 2
        if rows and abs(rows[-1][0] - y_mid) <= tolerance:
            rows[-1][1].append((x0, text))
        else:
            rows.append((y_mid, [(x0, text)]))
    return [(y_mid, sorted(entries, key=lambda entry: entry[0])) for y_mid, entries in rows]


def group_words_by_row(page: fitz.Page, tolerance: float = 3.0) -> List[List[Tuple[float, str]]]:
    rows = collect_word_rows(page, tolerance=tolerance)
    return [sorted(entries, key=lambda entry: entry[0]) for _, entries in rows]


def parse_answer_row(tokens: List[str]) -> List[str]:
    answers: List[str] = []
    for index, token in enumerate(tokens):
        normalized = normalize_ascii(token).replace(" ", "")
        if index == 0 and normalized.startswith("答案"):
            normalized = normalized[2:]
        if not normalized:
            continue
        normalized = normalized.replace("＃", "#")
        if re.fullmatch(r"[A-D#]+", normalized):
            answers.append(normalized)
    return answers


def parse_correction_clause(clause: str) -> Optional[str]:
    normalized = normalize_ascii(clause).replace(" ", "")
    if "更正為" in normalized:
        tail = normalized.split("更正為", 1)[1]
        choices = re.findall(r"[A-D]+", tail)
        if choices:
            return ",".join(dict.fromkeys(choices))
    if "不計分" in normalized:
        return "不計分"
    if any(keyword in normalized for keyword in ["一律給分", "均給分", "全部給分", "均予給分"]):
        choices = re.findall(r"[A-D]+", normalized)
        if choices:
            return ",".join(dict.fromkeys(choices))
        return "送分"
    choices = re.findall(r"[A-D]+", normalized)
    if choices:
        return ",".join(dict.fromkeys(choices))
    return None


def parse_answer_notes(full_text: str) -> Dict[int, str]:
    normalized_text = normalize_ascii(full_text)
    if "備" not in normalized_text or "註" not in normalized_text:
        return {}
    note_match = re.search(r"備\s*註[:：]\s*(.+)$", normalized_text, re.S)
    if not note_match:
        return {}
    note_text = note_match.group(1).strip()
    mappings: Dict[int, str] = {}
    for match in re.finditer(r"第(\d+)題(.+?)(?=第\d+題|$)", note_text):
        question_number = int(match.group(1))
        clause = match.group(2).strip("，。；; ")
        parsed = parse_correction_clause(clause)
        if parsed:
            mappings[question_number] = parsed
    return mappings


def parse_answer_pdf(path: Path, corrected: bool, errors: List[dict]) -> Tuple[Tuple[int, str, str], Dict[int, str]]:
    doc = fitz.open(path)
    year, exam_round, subject, class_name = parse_pdf_metadata(doc)
    answer_map: Dict[int, str] = {}

    if class_name != "醫事放射師":
        return (0, "", ""), answer_map

    if year is None or exam_round is None or not subject:
        errors.append(
            {
                "type": "parse_failed",
                "file": str(path),
                "message": "無法從答案 PDF 辨識年份、考試次別或科目",
            }
        )
        return (0, "", ""), answer_map

    for page in doc:
        rows = group_words_by_row(page)
        for index, row in enumerate(rows):
            tokens = [text for _, text in row]
            normalized_tokens = [normalize_ascii(token) for token in tokens]
            if not any(token in {"題號", "題序"} for token in normalized_tokens):
                continue
            numbers = [int(token) for token in normalized_tokens if token.isdigit()]
            if len(numbers) < 1 or index + 1 >= len(rows):
                continue
            answer_tokens = [text for _, text in rows[index + 1]]
            answers = parse_answer_row(answer_tokens)
            if len(answers) != len(numbers):
                continue
            for question_number, answer in zip(numbers, answers):
                answer_map[question_number] = answer

    if corrected:
        note_overrides = parse_answer_notes("\n".join(page.get_text("text") for page in doc))
        for question_number, answer in note_overrides.items():
            answer_map[question_number] = answer
        unresolved_hashes = [question_number for question_number, answer in answer_map.items() if answer == "#"]
        for question_number in unresolved_hashes:
            errors.append(
                {
                    "type": "unmatched_corrected_answer",
                    "year": year,
                    "exam_round": exam_round,
                    "subject": subject,
                    "question_number": question_number,
                    "file": str(path),
                    "message": "更正答案表中題號標示為 #，但備註沒有對應內容",
                }
            )
            del answer_map[question_number]

    return (year, exam_round, subject), answer_map


def parse_docx(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                parts.append("\t".join(values))
    return "\n".join(parts)


def parse_xlsx(path: Path) -> str:
    workbook = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if values:
                parts.append("\t".join(values))
    return "\n".join(parts)


def parse_text_with_fallback(path: Path) -> str:
    encodings = ["utf-8", "utf-8-sig", "big5", "cp950"]
    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except Exception:
            continue
    raise UnicodeDecodeError("unknown", b"", 0, 1, "無法解碼文字檔")


def merge_answers(
    questions: Dict[QUESTION_KEY, dict],
    original_answers: Dict[Tuple[int, str, str], Dict[int, str]],
    corrected_answers: Dict[Tuple[int, str, str], Dict[int, str]],
    errors: List[dict],
) -> List[dict]:
    merged_questions: List[dict] = []

    corrected_keys_seen = set()
    for key, question in sorted(questions.items()):
        year, exam_round, subject, question_number = key
        answer_context_key = (year, exam_round, subject)
        final_answer = None
        answer_source = None

        corrected_for_subject = corrected_answers.get(answer_context_key, {})
        original_for_subject = original_answers.get(answer_context_key, {})

        if question_number in corrected_for_subject:
            final_answer = corrected_for_subject[question_number]
            answer_source = "更正答案"
            corrected_keys_seen.add((answer_context_key, question_number))
        elif question_number in original_for_subject:
            final_answer = original_for_subject[question_number]
            answer_source = "原始答案"

        if final_answer is None:
            errors.append(
                {
                    "type": "missing_answer",
                    "year": year,
                    "exam_round": exam_round,
                    "subject": subject,
                    "question_number": question_number,
                    "file": question.get("_source_file", ""),
                    "message": "找不到此題對應答案",
                }
            )
            continue

        question["answer"] = final_answer
        question["answer_source"] = answer_source
        question.pop("_source_file", None)
        merged_questions.append(question)

    for answer_context_key, answer_map in corrected_answers.items():
        year, exam_round, subject = answer_context_key
        for question_number in answer_map:
            if (answer_context_key, question_number) not in corrected_keys_seen and (
                year,
                exam_round,
                subject,
                question_number,
            ) not in questions:
                errors.append(
                    {
                        "type": "unmatched_corrected_answer",
                        "year": year,
                        "exam_round": exam_round,
                        "subject": subject,
                        "question_number": question_number,
                        "file": "",
                        "message": "更正答案找不到對應題目",
                    }
                )

    return merged_questions


def validate_output(questions: List[dict], errors: List[dict]) -> None:
    required_fields = [
        "id",
        "year",
        "exam_round",
        "subject",
        "question_number",
        "question",
        "options",
        "answer",
        "answer_source",
        "explanation",
        "source",
        "images",
    ]
    seen_ids = set()
    for question in questions:
        for field in required_fields:
            if field not in question:
                errors.append(
                    {
                        "type": "parse_failed",
                        "file": question.get("id", ""),
                        "message": f"缺少必要欄位：{field}",
                    }
                )
        if list(question["options"].keys()) != ["A", "B", "C", "D"]:
            errors.append(
                {
                    "type": "parse_failed",
                    "file": question["id"],
                    "message": "options 順序不是 A、B、C、D",
                }
            )
        if question["id"] in seen_ids:
            errors.append(
                {
                    "type": "duplicated_question_id",
                    "file": question["id"],
                    "message": "偵測到重複題目 id",
                }
            )
        seen_ids.add(question["id"])


def build_summary_json(questions: List[dict], errors: List[dict]) -> dict:
    subjects = Counter(question["subject"] for question in questions)
    years = Counter(str(question["year"]) for question in questions)
    exam_rounds = Counter(question["exam_round"] for question in questions)
    answer_sources = Counter(question["answer_source"] for question in questions)
    image_questions = sum(1 for question in questions if question["images"])
    total_images = sum(len(question["images"]) for question in questions)
    unmatched_corrected_answers = sum(
        1 for error in errors if error.get("type") == "unmatched_corrected_answer"
    )
    return {
        "total_questions": len(questions),
        "subjects": dict(subjects),
        "years": dict(years),
        "exam_rounds": dict(exam_rounds),
        "answer_sources": dict(answer_sources),
        "image_questions": image_questions,
        "total_images": total_images,
        "errors": len(errors),
        "unmatched_corrected_answers": unmatched_corrected_answers,
    }


def write_json(path: Path, payload: object) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> int:
    if len(sys.argv) < 2:
        print("用法：python convert_exam_data_to_json.py questions_original_data")
        return 1

    input_dir = resolve_input_folder(sys.argv[1])
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    if ASSETS_DIR.exists():
        shutil.rmtree(ASSETS_DIR)
        ASSETS_DIR.mkdir(parents=True, exist_ok=True)

    errors: List[dict] = []
    initial_files = scan_input_folder(input_dir)
    zip_paths = [path for path in initial_files if detect_file_type(path) == "zip"]
    _, extracted_zip_count, extraction_errors = extract_zip_files(zip_paths)
    errors.extend(extraction_errors)

    all_files = sorted(
        [path for path in WORK_DIR.rglob("*") if path.is_file()]
        + [path for path in initial_files if detect_file_type(path) != "zip"]
    )

    question_files: List[Path] = []
    answer_files: List[Path] = []
    corrected_answer_files: List[Path] = []

    for path in all_files:
        file_type = detect_file_type(path)
        if file_type == "pdf":
            role = detect_document_role(path)
            if role == "question":
                question_files.append(path)
            elif role == "answer":
                answer_files.append(path)
            elif role == "corrected_answer":
                corrected_answer_files.append(path)
        elif file_type in {"docx", "xlsx", "text", "image"}:
            continue
        elif path.suffix.lower() == ".html" or path.name in {".DS_Store", "icon_link1.gif", "icon_link2.gif", "icon_link3.gif"}:
            continue
        else:
            errors.append(
                {
                    "type": "unsupported_file",
                    "file": str(path),
                    "message": "目前未處理的檔案格式",
                }
            )

    question_bank: Dict[QUESTION_KEY, dict] = {}
    for question_file in question_files:
        for question in parse_question_pdf(question_file, errors):
            key = (
                question["year"],
                question["exam_round"],
                question["subject"],
                question["question_number"],
            )
            if key in question_bank:
                errors.append(
                    {
                        "type": "duplicated_question_id",
                        "year": question["year"],
                        "exam_round": question["exam_round"],
                        "subject": question["subject"],
                        "question_number": question["question_number"],
                        "file": str(question_file),
                        "message": "題目重複",
                    }
                )
                continue
            question["_source_file"] = str(question_file)
            question_bank[key] = question

    original_answers: Dict[Tuple[int, str, str], Dict[int, str]] = {}
    corrected_answers: Dict[Tuple[int, str, str], Dict[int, str]] = {}

    for answer_file in answer_files:
        context_key, answer_map = parse_answer_pdf(answer_file, corrected=False, errors=errors)
        if context_key != (0, "", ""):
            original_answers[context_key] = answer_map

    for corrected_file in corrected_answer_files:
        context_key, answer_map = parse_answer_pdf(corrected_file, corrected=True, errors=errors)
        if context_key != (0, "", ""):
            corrected_answers[context_key] = answer_map

    output_questions = merge_answers(question_bank, original_answers, corrected_answers, errors)
    validate_output(output_questions, errors)
    summary = build_summary_json(output_questions, errors)

    write_json(OUTPUT_JSON, output_questions)
    write_json(ERRORS_JSON, errors)
    write_json(SUMMARY_JSON, summary)

    print(f"掃描到的檔案數: {len(initial_files)}")
    print(f"解壓縮的 zip 數: {extracted_zip_count}")
    print(f"偵測到的題目檔數: {len(question_files)}")
    print(f"偵測到的原始答案檔數: {len(answer_files)}")
    print(f"偵測到的更正答案檔數: {len(corrected_answer_files)}")
    print(f"成功輸出的題數: {summary['total_questions']}")
    print("各科題數:")
    for subject, count in summary["subjects"].items():
        print(f"  - {subject}: {count}")
    print(f"使用原始答案題數: {summary['answer_sources'].get('原始答案', 0)}")
    print(f"使用更正答案題數: {summary['answer_sources'].get('更正答案', 0)}")
    print(f"錯誤數: {summary['errors']}")
    print(f"output.json 路徑: {OUTPUT_JSON}")
    print(f"errors.json 路徑: {ERRORS_JSON}")
    print(f"summary.json 路徑: {SUMMARY_JSON}")
    print(f"assets/ 路徑: {ASSETS_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
